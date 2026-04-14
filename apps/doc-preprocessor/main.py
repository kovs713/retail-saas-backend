import json
import re
from enum import Enum
from io import BytesIO

import uvicorn
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from pypdf import PdfReader


class DocumentType(str, Enum):
    txt = "txt"
    md = "md"
    json = "json"
    docx = "docx"
    pdf = "pdf"


class TextPreprocessRequest(BaseModel):
    text: str
    remove_noise: bool = True
    normalize_whitespace: bool = True
    lowercase: bool = False


app = FastAPI(
    title="Document Preprocessor Service",
    description="Preprocess and convert documents between supported formats.",
    version="0.1.0",
)


def detect_type(filename: str | None, content_type: str | None) -> DocumentType:
    if filename and "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()
        if ext in DocumentType._value2member_map_:
            return DocumentType(ext)

    if content_type == "application/pdf":
        return DocumentType.pdf
    if content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    }:
        return DocumentType.docx
    if content_type == "application/json":
        return DocumentType.json
    if content_type in {"text/plain", "text/markdown"}:
        return DocumentType.txt

    raise HTTPException(
        status_code=400,
        detail="Cannot detect source type. Use known extension or content type.",
    )


def extract_text(content: bytes, source_type: DocumentType) -> str:
    if source_type in {DocumentType.txt, DocumentType.md}:
        return content.decode("utf-8", errors="replace")

    if source_type == DocumentType.json:
        payload = json.loads(content.decode("utf-8", errors="replace"))
        if isinstance(payload, dict) and "text" in payload:
            return str(payload["text"])
        return json.dumps(payload, ensure_ascii=False)

    if source_type == DocumentType.docx:
        doc = Document(BytesIO(content))
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    if source_type == DocumentType.pdf:
        reader = PdfReader(BytesIO(content))
        chunks: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            chunks.append(page_text)
        return "\n".join(chunks)

    raise HTTPException(
        status_code=400, detail=f"Unsupported source type: {source_type.value}"
    )


def preprocess_text(
    text: str,
    remove_noise: bool,
    normalize_whitespace: bool,
    lowercase: bool,
) -> str:
    cleaned = text

    if remove_noise:
        cleaned = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", cleaned)
        cleaned = re.sub(r"[\u200b-\u200f\ufeff]", "", cleaned)
        cleaned = re.sub(r"([!?.,])\1{2,}", r"\1", cleaned)

    if normalize_whitespace:
        cleaned = re.sub(r"\r\n?", "\n", cleaned)
        cleaned = re.sub(r"[ \t]+", " ", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
        cleaned = "\n".join(line.strip() for line in cleaned.split("\n")).strip()

    if lowercase:
        cleaned = cleaned.lower()

    return cleaned


def convert_text_to_type(
    text: str,
    target_type: DocumentType,
    source_name: str,
    source_type: DocumentType,
) -> tuple[bytes, str, str]:
    stem = (
        source_name.rsplit(".", 1)[0]
        if source_name and "." in source_name
        else "document"
    )
    out_name = f"{stem}.{target_type.value}"

    if target_type in {DocumentType.txt, DocumentType.md}:
        media = "text/plain; charset=utf-8"
        return text.encode("utf-8"), out_name, media

    if target_type == DocumentType.json:
        media = "application/json"
        payload = {
            "text": text,
            "meta": {
                "source_type": source_type.value,
                "target_type": target_type.value,
            },
        }
        return (
            json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"),
            out_name,
            media,
        )

    if target_type == DocumentType.docx:
        doc = Document()
        for line in text.split("\n"):
            doc.add_paragraph(line)
        buf = BytesIO()
        doc.save(buf)
        media = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        return buf.getvalue(), out_name, media

    raise HTTPException(
        status_code=400,
        detail=f"Target type {target_type.value} is not supported for output.",
    )


@app.get("/health")
async def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/preprocess-text")
async def preprocess_plain_text(payload: TextPreprocessRequest) -> dict[str, str]:
    cleaned = preprocess_text(
        text=payload.text,
        remove_noise=payload.remove_noise,
        normalize_whitespace=payload.normalize_whitespace,
        lowercase=payload.lowercase,
    )
    return {"text": cleaned}


@app.post("/preprocess-document")
async def preprocess_document(
    file: UploadFile = File(...),
    target_type: DocumentType = Form(DocumentType.txt),
    source_type: DocumentType | None = Form(None),
    remove_noise: bool = Form(True),
    normalize_whitespace: bool = Form(True),
    lowercase: bool = Form(False),
) -> StreamingResponse:
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Empty file.")

    src_type = source_type or detect_type(file.filename, file.content_type)
    text = extract_text(content=content, source_type=src_type)
    cleaned = preprocess_text(
        text=text,
        remove_noise=remove_noise,
        normalize_whitespace=normalize_whitespace,
        lowercase=lowercase,
    )
    payload, out_name, media_type = convert_text_to_type(
        text=cleaned,
        target_type=target_type,
        source_name=file.filename or "document",
        source_type=src_type,
    )

    return StreamingResponse(
        BytesIO(payload),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{out_name}"'},
    )


if __name__ == "__main__":
    uvicorn.run(app, port=1111)
